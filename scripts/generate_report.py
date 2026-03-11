#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
肾病患者健康生活方式建议报告生成脚本

根据患者信息和营养数据，生成 Word 格式的健康生活方式建议报告。

使用方法:
    python3 generate_report.py --name "张三" --age 55 --gender "男" ... --output report.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Dict, Any, Optional, List

# 尝试导入 docx 库
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx，运行：pip install python-docx")
    sys.exit(1)


def calculate_bmi(weight: float, height_cm: float) -> tuple:
    """计算 BMI 并返回数值和等级"""
    height_m = height_cm / 100
    bmi = weight / (height_m ** 2)

    # 中国标准
    if bmi < 18.5:
        level = "体重过低"
    elif bmi < 24:
        level = "正常范围"
    elif bmi < 28:
        level = "超重"
    else:
        level = "肥胖"

    return round(bmi, 1), level


def calculate_waist_risk(waist_cm: float, gender: str) -> str:
    """根据腰围判断中心性肥胖风险"""
    if gender in ['男', '男性', 'M']:
        if waist_cm >= 90:
            return "中心性肥胖"
        else:
            return "腰围正常"
    else:
        if waist_cm >= 85:
            return "中心性肥胖"
        else:
            return "腰围正常"


def parse_problems(problems_str: str) -> List[Dict[str, str]]:
    """解析核心问题字符串，返回问题列表"""
    problems = []

    if not problems_str:
        return problems

    # 支持多种格式
    lines = problems_str.replace('，', ',').replace(';', ',').replace('；', ',').split(',')

    for line in lines:
        line = line.strip()
        if line:
            # 尝试解析"问题：描述"格式
            if ':' in line or '：' in line:
                parts = line.replace('：', ':').split(':', 1)
                problems.append({
                    'name': parts[0].strip(),
                    'description': parts[1].strip() if len(parts) > 1 else '',
                    'goal': f"改善{parts[0].strip()}"
                })
            else:
                problems.append({
                    'name': line,
                    'description': '',
                    'goal': f"改善{line}"
                })

    return problems[:3]  # 最多 3 个问题


def create_basic_info_table(doc: Document, patient: Dict[str, Any]) -> None:
    """创建基本信息表格"""
    # 计算 BMI
    bmi, bmi_level = calculate_bmi(patient['weight'], patient['height'])
    waist_risk = calculate_waist_risk(patient.get('waist', 0), patient.get('gender', ''))

    # 6 列表格
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'

    # 第一行
    cells = table.rows[0].cells
    cells[0].text = "姓名"
    cells[1].text = patient.get('name', '')
    cells[2].text = "年龄（岁）"
    cells[3].text = str(patient.get('age', ''))
    cells[4].text = "性别"
    cells[5].text = patient.get('gender', '')

    # 第二行
    cells = table.rows[1].cells
    cells[0].text = "身高（cm）"
    cells[1].text = str(patient.get('height', ''))
    cells[2].text = "体重（kg）"
    cells[3].text = str(patient.get('weight', ''))
    cells[4].text = "腰围（cm）"
    cells[5].text = str(patient.get('waist', ''))

    # 第三行 - 生活习惯
    cells = table.rows[2].cells
    cells[0].text = "生活习惯"
    cells[1].text = patient.get('habits', '')
    for i in range(2, 6):
        cells[i].text = ""

    # 第四行 - 饮食现状
    cells = table.rows[3].cells
    cells[0].text = "饮食现状"
    cells[1].text = patient.get('diet_status', '')
    for i in range(2, 6):
        cells[i].text = ""

    # 第五行 - 既往史
    cells = table.rows[4].cells
    cells[0].text = "既往史"
    cells[1].text = patient.get('medical_history', '')
    for i in range(2, 6):
        cells[i].text = ""

    # 第六行 - 具体分析
    cells = table.rows[5].cells
    cells[0].text = "具体分析"
    analysis = f"1. BMI={bmi}，属于**{bmi_level}**范围，腰围{waist_risk}；"
    if patient.get('indicators'):
        analysis += f" 2. 检查指标：{patient.get('indicators', '')}。"
    cells[1].text = analysis
    cells[1].paragraphs[0].runs[0].bold = True
    for i in range(2, 6):
        cells[i].text = ""

    # 合并第六行的部分单元格
    table.rows[5].cells[1].merge(table.rows[5].cells[5])

    doc.add_paragraph()  # 空行


def create_problems_table(doc: Document, problems: List[Dict[str, str]]) -> None:
    """创建问题和管理目标表格"""
    table = doc.add_table(rows=len(problems) + 1, cols=3)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "核心问题"
    header_cells[1].text = "现状描述"
    header_cells[2].text = "管理目标"

    # 数据行
    for i, problem in enumerate(problems):
        cells = table.rows[i + 1].cells
        cells[0].text = problem.get('name', '')
        cells[1].text = problem.get('description', '')
        cells[2].text = problem.get('goal', '')

    doc.add_paragraph()  # 空行


def create_nutrition_table(doc: Document, nutrition_data: Dict[str, Any], patient: Dict[str, Any], period: Dict[str, str]) -> None:
    """创建营养干预方案表格"""
    # 计算营养需求
    weight = patient.get('weight', 60)
    height = patient.get('height', 170)
    age = patient.get('age', 50)
    gender = patient.get('gender', '男')

    # 基础代谢率 (Mifflin-St Jeor 公式)
    if gender in ['男', '男性', 'M']:
        bmr = 10 * weight + 6.25 * height - 5 * age + 5
    else:
        bmr = 10 * weight + 6.25 * height - 5 * age - 161

    # 每日能量需求（轻体力活动）
    energy_low = int(bmr * 1.2 * 0.85)  # 肾病患者的 85%
    energy_high = int(bmr * 1.2 * 0.95)

    # 蛋白质需求（肾病患者 0.6-0.8g/kg）
    protein_low = round(weight * 0.6, 1)
    protein_high = round(weight * 0.8, 1)

    # 碳水化合物（占总能量 50-55%）
    carb_low = int(energy_low * 0.50 / 4)
    carb_high = int(energy_high * 0.55 / 4)

    # 脂肪（占总能量 25-30%）
    fat_low = int(energy_low * 0.25 / 9)
    fat_high = int(energy_high * 0.30 / 9)

    table = doc.add_table(rows=13, cols=5)
    table.style = 'Table Grid'

    # 第一行 - 调理周期
    cells = table.rows[0].cells
    cells[0].text = f"营养调理周期：{period.get('start', '')}~{period.get('end', '')}"
    for i in range(1, 5):
        cells[i].text = ""
    cells[0].merge(cells[4])

    # 表头
    cells = table.rows[1].cells
    cells[0].text = "营养素"
    cells[1].text = "摄入量"
    cells[2].text = "参考依据"
    cells[3].text = ""
    cells[4].text = "备注"

    # 数据行
    nutrition_rows = [
        ("每日能量", f"{energy_low}-{energy_high}kcal/d", f"Mifflin-St Jeor 公式×活动系数", "", "根据肾功能调整"),
        ("蛋白质", f"{protein_low}-{protein_high}g/d", "0.6-0.8g/kg·d", "", "优质蛋白为主"),
        ("碳水化合物", f"{carb_low}-{carb_high}g/d", "占总能量 50-55%", "", "选择低 GI 食物"),
        ("脂肪", f"{fat_low}-{fat_high}g/d", "占总能量 25-30%", "", "限制饱和脂肪"),
        ("钠", "食盐<5g/d", "<2000mg 钠/d", "", "高血压患者<3g"),
        ("钾", "根据血钾水平", "3.5-5.5mmol/L", "正常范围", "高钾需限制"),
        ("磷", "<800mg/d", "CKD 患者限制", "0.8-1.5mmol/L", "避免高磷食物"),
        ("钙", "800-1000mg/d", "预防骨质疏松", "上限 2000mg", "配合维生素 D"),
        ("铁", "12-20mg/d", "预防贫血", "", "肾病易缺铁"),
        ("水", "根据尿量调整", "出入平衡", "", "水肿者需限水"),
    ]

    for i, row_data in enumerate(nutrition_rows):
        cells = table.rows[i + 2].cells
        for j, value in enumerate(row_data):
            cells[j].text = value

    # 餐次比行
    meal_ratio = nutrition_data.get('meal_ratio', {'breakfast': 30, 'lunch': 40, 'dinner': 30})
    cells = table.rows[12].cells
    cells[0].text = f"餐次比：早餐{meal_ratio.get('breakfast', 30)}%/午餐{meal_ratio.get('lunch', 40)}%/晚餐{meal_ratio.get('dinner', 30)}%"
    for i in range(1, 5):
        cells[i].text = ""
    cells[0].merge(cells[4])

    doc.add_paragraph()  # 空行


def create_exercise_table(doc: Document, patient: Dict[str, Any]) -> None:
    """创建运动方案表格（FITTVP 原则）"""
    # 引导语
    intro = doc.add_paragraph()
    intro.add_run("根据患者的身体状况和管理目标，制定以下个性化运动处方：").italic = True

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # 简化版运动建议（实际应用中可根据患者情况调整）
    exercise_data = [
        ("频率 (F)", "每周 3-5 次，隔天进行"),
        ("强度 (I)", "中等强度，运动时微微出汗，能说话但不能唱歌"),
        ("时间 (T)", "每次 30-45 分钟，包括热身 5 分钟和整理运动 5 分钟"),
        ("类型 (T)", "有氧运动为主：快走、慢跑、游泳、骑自行车；配合抗阻训练"),
        ("时机", "餐后 1 小时进行，避免空腹运动；下午或傍晚为宜"),
    ]

    for i, (key, value) in enumerate(exercise_data):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = value

    doc.add_paragraph()  # 空行


def create_psychological_section(doc: Document) -> None:
    """创建心理与睡眠建议部分"""
    suggestions = [
        ("压力管理", "每天进行 15-20 分钟的深呼吸或冥想练习，帮助缓解压力。可尝试渐进性肌肉放松法。"),
        ("情绪调节", "保持积极乐观的心态，多与家人朋友交流，必要时寻求专业心理咨询。"),
        ("睡眠卫生", "保持规律作息，每晚 10 点前入睡，保证 7-8 小时睡眠。睡前避免使用电子设备。"),
        ("兴趣爱好", "培养兴趣爱好如听音乐、养花、书法等，丰富生活内容，提升生活质量。"),
    ]

    for i, (title, content) in enumerate(suggestions, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title}：").bold = True
        p.add_run(content)

    doc.add_paragraph()  # 空行


def create_summary_section(doc: Document, patient: Dict[str, Any], manager: Dict[str, str]) -> None:
    """创建总结和关键提醒部分"""
    # 总结性语句
    name = patient.get('name', '您')
    summary_intro = doc.add_paragraph()
    summary_intro.add_run(f"{name}先生/女士，健康管理是一个循序渐进的过程，需要您的坚持和配合。以下是关键提醒：").bold = True

    reminders = [
        ("定期监测", "按时测量血压、血糖、体重等指标，记录变化趋势，复诊时带给医生参考。"),
        ("遵医嘱用药", "按时按量服用药物，不随意增减剂量或停药。如有不适及时就医。"),
        ("饮食控制", "严格遵循营养建议，低盐低脂优质蛋白饮食，戒烟限酒，保持健康生活方式。"),
    ]

    for i, (title, content) in enumerate(reminders, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {title}：").bold = True
        p.add_run(content)

    # 署名
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"调理健管师：{manager.get('name', '')}\n").bold = True
    p.add_run(f"干预日期：{manager.get('date', '')}")


def generate_report(
    patient: Dict[str, Any],
    nutrition_data: Dict[str, Any],
    period: Dict[str, str],
    manager: Dict[str, str],
    output_path: str,
    organization: str = ""
) -> str:
    """生成完整的 Word 报告"""
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 报告头部
    if organization:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(f"{organization}\n").bold = True
        title.add_run("肾病患者健康生活方式建议报告").bold = True
    else:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("肾病患者健康生活方式建议报告").bold = True

    # 患者信息
    info = doc.add_paragraph()
    info.add_run(f"患者姓名：{patient.get('name', '')}    ")
    info.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}")

    doc.add_paragraph()  # 空行

    # 1. 基本信息
    h1 = doc.add_heading("1. 基本信息", level=1)
    create_basic_info_table(doc, patient)

    # 2. 问题和管理目标
    h2 = doc.add_heading("2. 问题和管理目标", level=1)
    problems = parse_problems(patient.get('problems', ''))
    if not problems:
        problems = [
            {'name': '体重管理', 'description': patient.get('diet_status', ''), 'goal': '控制体重在正常范围'},
            {'name': '营养调理', 'description': '饮食结构需优化', 'goal': '均衡营养，改善饮食习惯'},
            {'name': '慢病管理', 'description': patient.get('medical_history', ''), 'goal': '控制病情发展，提高生活质量'},
        ]
    create_problems_table(doc, problems)

    # 3. 营养干预方案
    h3 = doc.add_heading("3. 营养干预方案", level=1)

    h3_1 = doc.add_heading("3.1 营养干预原则", level=2)
    principles = doc.add_paragraph()
    principles.add_run("• 优质低蛋白\n")
    principles.add_run("• 低盐低钠\n")
    principles.add_run("• 充足能量\n")
    principles.add_run("• 适量矿物质\n")
    principles.add_run("• 均衡营养")

    h3_2 = doc.add_heading("3.2 营养干预具体方案", level=2)
    create_nutrition_table(doc, nutrition_data, patient, period)

    h3_3 = doc.add_heading("3.3 营养干预详细计划（第一周）", level=2)
    week1_table = doc.add_table(rows=4, cols=2)
    week1_table.style = 'Table Grid'
    week1_data = [
        ("目的", "适应新的饮食结构，建立良好饮食习惯"),
        ("饮食方案", "按照上述营养方案执行，记录每日饮食和身体反应"),
        ("饮食注意事项", "1.少食多餐 2.细嚼慢咽 3.避免暴饮暴食 4.限制加工食品 5.多吃新鲜蔬果"),
    ]
    for i, (key, value) in enumerate(week1_data):
        cells = week1_table.rows[i].cells
        cells[0].text = key
        cells[1].text = value

    doc.add_paragraph()  # 空行

    # 4. 运动方案
    h4 = doc.add_heading("4. 运动方案", level=1)
    create_exercise_table(doc, patient)

    # 5. 心理与睡眠建议
    h5 = doc.add_heading("5. 心理与睡眠建议", level=1)
    create_psychological_section(doc)

    # 总结和关键提醒
    h6 = doc.add_heading("总结和关键提醒", level=1)
    create_summary_section(doc, patient, manager)

    # 保存文档
    doc.save(output_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(description='生成肾病患者健康生活方式建议报告')

    # 患者基本信息
    parser.add_argument('--name', required=True, help='患者姓名')
    parser.add_argument('--age', type=int, required=True, help='年龄')
    parser.add_argument('--gender', required=True, help='性别')
    parser.add_argument('--height', type=float, required=True, help='身高 (cm)')
    parser.add_argument('--weight', type=float, required=True, help='体重 (kg)')
    parser.add_argument('--waist', type=float, default=0, help='腰围 (cm)')
    parser.add_argument('--habits', default='', help='生活习惯')
    parser.add_argument('--diet-status', default='', help='饮食现状')
    parser.add_argument('--medical-history', default='', help='既往病史')
    parser.add_argument('--indicators', default='', help='检查指标')
    parser.add_argument('--problems', default='', help='核心问题')

    # 营养报告
    parser.add_argument('--nutrition-report', default='', help='营养报告文件路径')
    parser.add_argument('--nutrition-data', default='', help='营养数据 JSON 字符串或文件路径')

    # 干预周期
    parser.add_argument('--start-date', default='', help='开始日期')
    parser.add_argument('--end-date', default='', help='结束日期')

    # 健管师信息
    parser.add_argument('--manager-name', default='', help='健管师姓名')
    parser.add_argument('--manager-date', default='', help='干预日期')

    # 其他
    parser.add_argument('--organization', default='', help='机构名称')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')

    args = parser.parse_args()

    # 构建患者信息字典
    patient = {
        'name': args.name,
        'age': args.age,
        'gender': args.gender,
        'height': args.height,
        'weight': args.weight,
        'waist': args.waist,
        'habits': args.habits,
        'diet_status': args.diet_status,
        'medical_history': args.medical_history,
        'indicators': args.indicators,
        'problems': args.problems,
    }

    # 处理营养数据
    nutrition_data = {}
    if args.nutrition_report:
        # 从营养报告提取
        extract_script = os.path.join(os.path.dirname(__file__), 'extract_nutrition_data.py')
        if os.path.exists(extract_script):
            import subprocess
            result = subprocess.run(
                ['python3', extract_script, '-i', args.nutrition_report, '-o', '/tmp/nutrition_data.json'],
                capture_output=True, text=True
            )
            if os.path.exists('/tmp/nutrition_data.json'):
                with open('/tmp/nutrition_data.json', 'r', encoding='utf-8') as f:
                    nutrition_data = json.load(f)
    elif args.nutrition_data:
        # 从 JSON 字符串或文件加载
        if os.path.exists(args.nutrition_data):
            with open(args.nutrition_data, 'r', encoding='utf-8') as f:
                nutrition_data = json.load(f)
        else:
            nutrition_data = json.loads(args.nutrition_data)

    # 干预周期
    period = {
        'start': args.start_date or datetime.now().strftime('%Y-%m-%d'),
        'end': args.end_date or '',
    }

    # 健管师信息
    manager = {
        'name': args.manager_name or '健管师',
        'date': args.manager_date or datetime.now().strftime('%Y-%m-%d'),
    }

    # 生成报告
    output_path = generate_report(
        patient=patient,
        nutrition_data=nutrition_data,
        period=period,
        manager=manager,
        output_path=args.output,
        organization=args.organization,
    )

    print(f"报告已生成：{output_path}")


if __name__ == '__main__':
    main()
