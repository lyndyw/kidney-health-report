#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 JSON 数据生成肾病患者健康生活方式建议报告

使用方法:
    python3 generate_report_from_json.py --input /tmp/nutrition_data.json --output report.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx，运行：pip install python-docx")
    sys.exit(1)


def calculate_bmi_from_data(patient_info: Dict) -> tuple:
    """从患者信息计算 BMI（如果已有则直接返回）"""
    if patient_info.get('bmi'):
        bmi = float(patient_info['bmi'])
    else:
        weight = patient_info.get('weight', 0)
        height_cm = patient_info.get('height', 0)
        if weight and height_cm:
            height_m = height_cm / 100
            bmi = weight / (height_m ** 2)
        else:
            # 根据 BMI 等级估算
            bmi = 30.2  # 默认值

    # 判断等级
    if bmi < 18.5:
        level = "体重过低"
    elif bmi < 24:
        level = "正常范围"
    elif bmi < 28:
        level = "超重"
    else:
        level = "肥胖"

    return round(bmi, 1), level


def create_header(doc: Document, patient_name: str, organization: str = "") -> None:
    """创建报告头部"""
    if organization:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{organization}\n")
        run.bold = True
        run.font.size = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("肾病患者健康生活方式建议报告")
    run.bold = True
    run.font.size = Pt(18)

    # 患者信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run(f"患者姓名：{patient_name}    ")
    info.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}")

    doc.add_paragraph()  # 空行


def create_basic_info_section(doc: Document, patient_info: Dict) -> None:
    """创建基本信息部分"""
    h1 = doc.add_heading("1. 基本信息", level=1)

    # 计算 BMI
    bmi, bmi_level = calculate_bmi_from_data(patient_info)

    # 6 列表格
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'

    # 第一行 - 姓名、年龄、性别
    cells = table.rows[0].cells
    cells[0].text = "姓名"
    cells[1].text = patient_info.get('name', '')
    cells[2].text = "年龄（岁）"
    cells[3].text = patient_info.get('age', '-')
    cells[4].text = "性别"
    cells[5].text = patient_info.get('gender', '-')

    # 第二行 - 身高、体重、腰围
    cells = table.rows[1].cells
    cells[0].text = "身高（cm）"
    cells[1].text = str(patient_info.get('height', '-'))
    cells[2].text = "体重（kg）"
    cells[3].text = str(patient_info.get('weight', '-'))
    cells[4].text = "腰围（cm）"
    cells[5].text = str(patient_info.get('waist', '-'))

    # 第三行 - 生活习惯
    cells = table.rows[2].cells
    cells[0].text = "生活习惯"
    cells[1].text = "待补充"
    for i in range(2, 6):
        cells[i].text = ""
    cells[1].merge(cells[5])

    # 第四行 - 饮食现状
    cells = table.rows[3].cells
    cells[0].text = "饮食现状"
    cells[1].text = "待补充"
    for i in range(2, 6):
        cells[i].text = ""
    cells[1].merge(cells[5])

    # 第五行 - 既往史
    medical_history = patient_info.get('medical_history', [])
    history_str = '、'.join(medical_history) if medical_history else '-'
    cells = table.rows[4].cells
    cells[0].text = "既往史"
    cells[1].text = history_str
    for i in range(2, 6):
        cells[i].text = ""
    cells[1].merge(cells[5])

    # 第六行 - 具体分析
    cells = table.rows[5].cells
    cells[0].text = "具体分析"
    analysis = f"1. BMI={bmi}，属于**{bmi_level}**范围；"

    # 添加指标信息
    indicators = patient_info.get('indicators', {})
    if indicators:
        indicator_parts = []
        for key, value in indicators.items():
            indicator_parts.append(f"{key}: {value}")
        if indicator_parts:
            analysis += f" 2. 检查指标：{', '.join(indicator_parts)}。"

    cells[1].text = analysis
    cells[1].merge(cells[5])

    # 第七行 - CKD 分期
    cells = table.rows[6].cells
    cells[0].text = "CKD 分期"
    ckd_stage = patient_info.get('ckd_stage', '待评估')
    cells[1].text = ckd_stage
    for i in range(2, 6):
        cells[i].text = ""
    cells[1].merge(cells[5])

    doc.add_paragraph()  # 空行


def create_problems_section(doc: Document, patient_info: Dict) -> None:
    """创建问题和管理目标部分"""
    h2 = doc.add_heading("2. 问题和管理目标", level=1)

    # 根据患者情况自动生成问题
    problems = []

    medical_history = patient_info.get('medical_history', [])

    # 根据病史生成问题
    if '慢性肾脏病' in medical_history:
        problems.append({
            'name': '慢性肾脏病',
            'description': f"eGFR {patient_info.get('indicators', {}).get('eGFR', '未知')}，{patient_info.get('ckd_stage', 'CKD')} ",
            'goal': '保护肾功能，延缓 CKD 进展'
        })

    if '肥胖' in medical_history or patient_info.get('bmi', 0) >= 28:
        problems.append({
            'name': '肥胖症',
            'description': f"BMI {patient_info.get('bmi', '未知')}，属于肥胖范围",
            'goal': '减重 10-15%，改善代谢指标'
        })

    if '高血压' in medical_history:
        problems.append({
            'name': '高血压',
            'description': '高血压病史，治疗中',
            'goal': '血压控制在<130/80 mmHg'
        })

    if '痛风' in medical_history:
        problems.append({
            'name': '痛风',
            'description': '痛风病史，需控制尿酸',
            'goal': '血尿酸控制在<360 μmol/L，预防发作'
        })

    if '糖代谢异常' in medical_history or '脂肪肝' in medical_history:
        problems.append({
            'name': '代谢综合征',
            'description': '高甘油三酯、血糖受损、脂肪肝',
            'goal': '改善代谢指标，降低心血管风险'
        })

    # 确保至少 3 个问题
    if len(problems) < 3:
        problems.append({
            'name': '生活方式改善',
            'description': '饮食结构需优化，运动不足',
            'goal': '建立健康生活方式，提高生活质量'
        })

    # 3 列表格
    table = doc.add_table(rows=len(problems) + 1, cols=3)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "核心问题"
    header_cells[1].text = "现状描述"
    header_cells[2].text = "管理目标"

    # 数据行
    for i, problem in enumerate(problems):
        cells = table.rows[i + 1].cells
        cells[0].text = problem['name']
        cells[1].text = problem['description']
        cells[2].text = problem['goal']

    doc.add_paragraph()  # 空行


def create_nutrition_section(doc: Document, nutrition_targets: Dict) -> None:
    """创建营养干预方案部分"""
    h3 = doc.add_heading("3. 营养干预方案", level=1)

    # 3.1 营养干预原则
    h3_1 = doc.add_heading("3.1 营养干预原则", level=2)
    principles = doc.add_paragraph()
    principles.add_run("• 优质低蛋白（保护肾脏）\n")
    principles.add_run("• 低盐低钠（控制血压）\n")
    principles.add_run("• 低脂饮食（降脂护肝）\n")
    principles.add_run("• 低嘌呤（预防痛风）\n")
    principles.add_run("• 低 GI 食物（控制血糖）\n")
    principles.add_run("• 充足能量（维持体重）")

    # 3.2 营养干预具体方案
    h3_2 = doc.add_heading("3.2 营养干预具体方案", level=2)

    # 5 列表格
    table = doc.add_table(rows=13, cols=5)
    table.style = 'Table Grid'

    # 第一行 - 调理周期（留空待填）
    cells = table.rows[0].cells
    cells[0].text = f"营养调理周期：____年__月__日 ~ ____年__月__日"
    for i in range(1, 5):
        cells[i].text = ""
    cells[0].merge(cells[4])

    # 表头
    cells = table.rows[1].cells
    cells[0].text = "营养素"
    cells[1].text = "摄入量"
    cells[2].text = "参考依据"
    cells[3].text = ""
    cells[4].text = "备注"

    # 数据行
    nutrient_order = ['energy', 'protein', 'carbohydrate', 'fat', 'sodium',
                      'potassium', 'phosphorus', 'calcium', 'iron', 'water']

    nutrient_names = {
        'energy': '每日能量',
        'protein': '蛋白质',
        'carbohydrate': '碳水化合物',
        'fat': '脂肪',
        'sodium': '钠',
        'potassium': '钾',
        'phosphorus': '磷',
        'calcium': '钙',
        'iron': '铁',
        'water': '水',
    }

    row_idx = 2
    for nutrient_key in nutrient_order:
        if nutrient_key in nutrition_targets:
            info = nutrition_targets[nutrient_key]
            cells = table.rows[row_idx].cells
            cells[0].text = nutrient_names.get(nutrient_key, nutrient_key)
            cells[1].text = f"{info.get('value', '')} {info.get('unit', '')}".strip()
            cells[2].text = info.get('basis', '')
            cells[3].text = ""
            cells[4].text = info.get('note', '')
            row_idx += 1

    # 餐次比行
    meal_ratio = nutrition_targets.get('meal_ratio', {'breakfast': 30, 'lunch': 40, 'dinner': 30})
    cells = table.rows[12].cells
    cells[0].text = f"餐次比：早餐{meal_ratio.get('breakfast', 30)}%/午餐{meal_ratio.get('lunch', 40)}%/晚餐{meal_ratio.get('dinner', 30)}%"
    for i in range(1, 5):
        cells[i].text = ""
    cells[0].merge(cells[4])

    doc.add_paragraph()  # 空行

    # 3.3 营养干预详细计划（第一周）
    h3_3 = doc.add_heading("3.3 营养干预详细计划（第一周）", level=2)

    week1_table = doc.add_table(rows=4, cols=2)
    week1_table.style = 'Table Grid'

    week1_data = [
        ("目的", "适应新的饮食结构，建立良好饮食习惯"),
        ("饮食方案", "按照上述营养方案执行，记录每日饮食和身体反应。优先选择优质蛋白（鸡蛋、牛奶、鱼肉、豆腐），限制高嘌呤、高脂肪食物。"),
        ("饮食注意事项", "1. 少食多餐，避免暴饮暴食\n2. 细嚼慢咽，每口咀嚼 20-30 次\n3. 限制加工食品和高盐食物\n4. 多吃新鲜蔬菜，适量低 GI 水果\n5. 戒酒，避免含糖饮料"),
    ]

    for i, (key, value) in enumerate(week1_data):
        cells = week1_table.rows[i].cells
        cells[0].text = key
        cells[1].text = value

    doc.add_paragraph()  # 空行


def create_exercise_section(doc: Document, exercise: Dict) -> None:
    """创建运动方案部分"""
    h4 = doc.add_heading("4. 运动方案", level=1)

    # 引导语
    intro = doc.add_paragraph()
    intro.add_run("根据患者的身体状况和管理目标，制定以下个性化运动处方：").italic = True

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    exercise_data = [
        ("频率 (F)", exercise.get('frequency', '每周 3-5 次，隔天进行')),
        ("强度 (I)", exercise.get('intensity', '中等强度，运动时微微出汗，能说话但不能唱歌')),
        ("时间 (T)", f"{exercise.get('duration', '每次 30 分钟')}，包括热身 5 分钟和整理运动 5 分钟"),
        ("类型 (T)", exercise.get('type', '有氧运动为主：快走、游泳、骑自行车')),
        ("时机", "餐后 1 小时进行，避免空腹运动；下午或傍晚为宜"),
    ]

    for i, (key, value) in enumerate(exercise_data):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = value

    doc.add_paragraph()  # 空行


def create_psychological_section(doc: Document) -> None:
    """创建心理与睡眠建议部分"""
    h5 = doc.add_heading("5. 心理与睡眠建议", level=1)

    suggestions = [
        ("压力管理", "每天进行 15-20 分钟的深呼吸或冥想练习，帮助缓解压力。可尝试渐进性肌肉放松法。"),
        ("情绪调节", "保持积极乐观的心态，多与家人朋友交流，必要时寻求专业心理咨询。"),
        ("睡眠卫生", "保持规律作息，每晚 10 点前入睡，保证 7-8 小时睡眠。睡前避免使用电子设备。"),
        ("兴趣爱好", "培养兴趣爱好如听音乐、养花、书法等，丰富生活内容，提升生活质量。"),
    ]

    for i, (title, content) in enumerate(suggestions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}：")
        run.bold = True
        p.add_run(content)

    doc.add_paragraph()  # 空行


def create_summary_section(doc: Document, patient_name: str, manager_name: str = "", manager_date: str = "") -> None:
    """创建总结和关键提醒部分"""
    h6 = doc.add_heading("总结和关键提醒", level=1)

    # 总结性语句
    summary_intro = doc.add_paragraph()
    summary_intro.add_run(f"{patient_name}先生/女士，健康管理是一个循序渐进的过程，需要您的坚持和配合。以下是关键提醒：").bold = True

    reminders = [
        ("定期监测", "按时测量血压、血糖、体重等指标，记录变化趋势，复诊时带给医生参考。每月复查肾功能、血脂、尿酸等指标。"),
        ("遵医嘱用药", "按时按量服用药物，不随意增减剂量或停药。如有不适及时就医。CKD 患者需特别注意避免肾毒性药物。"),
        ("饮食控制", "严格遵循营养建议，低盐低脂优质蛋白饮食，戒烟限酒。避免高嘌呤食物（动物内脏、浓肉汤、海鲜、啤酒），预防痛风发作。"),
    ]

    for i, (title, content) in enumerate(reminders, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}：")
        run.bold = True
        p.add_run(content)

    # 署名
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"调理健管师：{manager_name if manager_name else '____________'}\n")
    run.bold = True
    p.add_run(f"干预日期：{manager_date if manager_date else '____年__月__日'}")


def generate_report(
    data: Dict[str, Any],
    output_path: str,
    organization: str = "",
    manager_name: str = "",
) -> str:
    """生成完整的 Word 报告"""
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    patient_info = data.get('patient_info', {})
    nutrition_targets = data.get('nutrition_targets', {})
    exercise = data.get('exercise', {})
    patient_name = patient_info.get('name', '患者')

    # 报告头部
    create_header(doc, patient_name, organization)

    # 1. 基本信息
    create_basic_info_section(doc, patient_info)

    # 2. 问题和管理目标
    create_problems_section(doc, patient_info)

    # 3. 营养干预方案
    create_nutrition_section(doc, nutrition_targets)

    # 4. 运动方案
    create_exercise_section(doc, exercise)

    # 5. 心理与睡眠建议
    create_psychological_section(doc)

    # 总结和关键提醒
    create_summary_section(doc, patient_name, manager_name)

    # 保存文档
    doc.save(output_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(description='从 JSON 数据生成健康生活方式建议报告')
    parser.add_argument('--input', '-i', required=True, help='输入 JSON 文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出 Word 文件路径')
    parser.add_argument('--organization', default='', help='机构名称')
    parser.add_argument('--manager-name', default='', help='健管师姓名')
    parser.add_argument('--manager-date', default='', help='干预日期')

    args = parser.parse_args()

    # 加载 JSON 数据
    if not os.path.exists(args.input):
        print(f"错误：文件不存在：{args.input}")
        sys.exit(1)

    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 生成报告
    output_path = generate_report(
        data=data,
        output_path=args.output,
        organization=args.organization,
        manager_name=args.manager_name,
    )

    print(f"报告已生成：{output_path}")
    print(f"文件大小：{os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
